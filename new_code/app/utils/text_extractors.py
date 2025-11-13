# from typing import Tuple
# import os

# def extract_text_from_bytes(filename: str, data: bytes) -> Tuple[str, str]:
#     """
#     Returns (text, title). For demo we handle plain text & pdf minimal.
#     Plug in pypdf/docx2txt etc as you expand.
#     """
#     name = os.path.basename(filename)
#     lower = name.lower()
#     if lower.endswith(".txt"):
#         return data.decode("utf-8", errors="ignore"), name
#     if lower.endswith(".pdf"):
#         try:
#             from pypdf import PdfReader
#             import io
#             reader = PdfReader(io.BytesIO(data))
#             pages = [p.extract_text() or "" for p in reader.pages]
#             return "\n\n".join(pages), name
#         except Exception:
#             return "", name
#     # default: try utf-8
#     return data.decode("utf-8", errors="ignore"), name


# utils/text_extractors.py
from typing import Optional

from fastapi import HTTPException
# from Rag import *
import os
from app.Rag.ai import loader
from io import BytesIO
from langchain_core.documents import Document
def extract_full_text(pages_data):
    """
    Combine page_content from all pages into a single text string.
    
    Args:
        pages_data (list): List of dictionaries (each with 'page_content' key)
    
    Returns:
        str: Combined text from all pages
    """
    if not isinstance(pages_data, list):
        raise ValueError("Input must be a list of page data objects.")
    
    res=""
    for item in pages_data:
        res+=" "+item.page_content

    return res


def extract_text(file_bytes: bytes, filename: str, mimetype: Optional[str]) -> str:
    """
    Extracts text from common file types: PDF, DOCX, TXT.
    Falls back to UTF-8 decode if extractors fail.
    """
    name = (filename or "").lower()
    mt = (mimetype or "").lower()
    docs=[]
    try:
        # PDF
        if name.endswith(".pdf") or "pdf" in mt:
            try:
                import fitz  # PyMuPDF (best)
                from io import BytesIO
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                
                for i,p in enumerate(doc):
                    docs.append(Document(page_content=p.get_text("text"),metadata={"pages":i+1,"filename":filename}))
                text=extract_full_text(docs)
                # print("docs",docs)
                return text,docs
                # return "\n".join(p.get_text("text") for p in doc) ,
            except Exception:
                try:
                    from pypdf import PdfReader
                    from io import BytesIO
                    reader = PdfReader(BytesIO(file_bytes))
                    for i,p in enumerate(reader.pages):
                        docs.append(Document(page_content=p.extract_text(),metadata={"pages":i+1}))
                    text=extract_full_text(docs)
                    return text,docs

                    # return "\n".join((p.extract_text() or "") for p in reader.pages)
                except Exception:
                    pass

        # DOCX
        if name.endswith(".docx") or "officedocument.wordprocessingml.document" in mt:
            try:
                import docx
                from io import BytesIO
                d = docx.Document(BytesIO(file_bytes))
                for i,p in enumerate(doc):
                    docs.append(Document(page_content=p.get_text("text"),metadata={"pages":i+1,"filename":filename}))
                text=extract_full_text(docs)
                return text,docs
                # return "\n".join(p.text for p in d.paragraphs)
            except Exception:
                pass

        # Plain text-like
        return file_bytes.decode("utf-8", errors="ignore")

    except Exception:
        return file_bytes.decode("utf-8", errors="ignore")

