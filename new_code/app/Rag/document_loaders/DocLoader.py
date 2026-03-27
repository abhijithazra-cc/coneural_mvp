
# from app.Rag.abstractions.Idocloader import Idocloader

# from io import BytesIO
# from langchain_core.documents import Document

# import docx
# from io import BytesIO
# class DocLoader(Idocloader):
#      def __init__(self):
#           self.documents=None

#      def load_document(self,file:str | bytes,filename:str,id=None):
                
#                 if type(file)==str:
#                     print("file type","str")
#                     docs=[]

#                     doc = docx.Document(BytesIO(file))
#                     for i,p in enumerate(doc):
#                         docs.append(Document(page_content=p.get_text("text"),metadata={"pages":i+1,"filename":filename}))

#                     self.documents=docs
#                 else:
#                     print("file type","bytes")
#                     docs=[]
#                     doc = docx.Document(file)
#                     for i,p in enumerate(doc):
#                         docs.append(Document(page_content=p.get_text("text"),id=id,metadata={"pages":i+1,"filename":filename}))

#                     self.documents=docs
#                     print("my docs",docs)

#      def get_document(self):
#           return self.documents
#      def get_full_content(self):
#           res=""
#           for item in self.documents:
#               res+=" "+item.page_content

#           return res






from app.Rag.abstractions.Idocloader import Idocloader
from io import BytesIO
from langchain_core.documents import Document
import docx


class DocLoader(Idocloader):
    def __init__(self):
        self.documents = None

    def load_document(self, file: str | bytes, filename: str, id=None):
        docs = []

        if isinstance(file, str):
            # file is a filesystem path — pass directly
            print("file type", "str")
            doc = docx.Document(file)
        elif isinstance(file, bytes):
            # file is raw bytes — wrap in BytesIO so docx gets a seekable stream
            print("file type", "bytes")
            doc = docx.Document(BytesIO(file))
        else:
            raise TypeError(f"Unsupported file type: {type(file)}")

        for i, p in enumerate(doc.paragraphs):   # ← must iterate .paragraphs
            if p.text.strip():                    # skip empty paragraphs
                docs.append(Document(
                    page_content=p.text,          # ← .text, not .get_text()
                    id=id,
                    metadata={"pages": i + 1, "filename": filename}
                ))

        self.documents = docs
        print("my docs", docs)

    def get_document(self):
        return self.documents

    def get_full_content(self):
        res = ""
        for item in self.documents:
            res += " " + item.page_content
        return res